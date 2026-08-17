# -*- coding: utf-8 -*-
"""从 docs/05-商业计划书.md 生成《商业计划书-杏坛智备-V1.2.docx》。

设计令牌（延续项目品牌）：
- 页面：A4，上下 2.2cm / 左右 2.4cm
- 标题：微软雅黑加粗，品牌绿 #2E6B4F（一级/二级带下边框）
- 正文：宋体 11pt，深灰 #333333，1.3 倍行距
- 表格：品牌绿表头白字、浅绿斑马纹、显式 DXA 几何、跨页重复表头
- 列表：真实 Word 编号定义（编号列表每段重新从 1 开始）
"""
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[2]
MD = ROOT / "docs" / "05-商业计划书.md"
OUT = ROOT / "商业计划书-杏坛智备-V1.2.docx"

BRAND_GREEN = "2E6B4F"
GOLD = "E3A03C"
DARK = RGBColor(0x33, 0x33, 0x33)
GRAY = RGBColor(0x6B, 0x6B, 0x6B)
LIGHT_GREEN = "EAF3EE"
USABLE_DXA = 9500  # A4 21cm - 左右 2.4cm*2 ≈ 16.2cm ≈ 9185 DXA，留 315 余量


def set_run_font(run, size=11, bold=False, color=DARK, east="宋体", ascii_font="Calibri"):
    run.font.name = ascii_font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color


def add_runs(paragraph, text, base_size=11, bold=False, color=DARK):
    """支持 **加粗**、`代码` 内联标记。"""
    tokens = re.split(r"(\*\*.+?\*\*|`[^`]+`)", text)
    for tok in tokens:
        if not tok:
            continue
        if tok.startswith("**") and tok.endswith("**"):
            run = paragraph.add_run(tok[2:-2])
            set_run_font(run, base_size, True, color)
        elif tok.startswith("`") and tok.endswith("`"):
            run = paragraph.add_run(tok[1:-1])
            set_run_font(run, base_size - 1, bold, RGBColor(0xC7, 0x25, 0x4E), east="Consolas", ascii_font="Consolas")
        else:
            run = paragraph.add_run(tok)
            set_run_font(run, base_size, bold, color)


def add_heading(doc, text, level):
    sizes = {1: 18, 2: 15, 3: 13, 4: 12}
    p = doc.add_paragraph(style=f"Heading {min(level, 4)}")
    p.paragraph_format.space_before = Pt(16 if level <= 2 else 11)
    p.paragraph_format.space_after = Pt(7)
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    set_run_font(run, sizes.get(level, 12), True, RGBColor.from_string(BRAND_GREEN) if level <= 2 else DARK,
                 east="微软雅黑", ascii_font="Calibri")
    if level <= 2:
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "8")
        bottom.set(qn("w:space"), "3")
        bottom.set(qn("w:color"), BRAND_GREEN)
        pBdr.append(bottom)
        pPr.append(pBdr)
    return p


def add_para(doc, text, size=11):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.3
    add_runs(p, text, size)
    return p


def _numbering_part(doc):
    return doc.part.numbering_part.element


def _new_list_numbering(doc, numbered):
    numbering = _numbering_part(doc)
    abs_ids = [int(a.get(qn("w:abstractNumId"))) for a in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(n.get(qn("w:numId"))) for n in numbering.findall(qn("w:num"))]
    abs_id = (max(abs_ids) + 1) if abs_ids else 1
    num_id = (max(num_ids) + 1) if num_ids else 1

    an = OxmlElement("w:abstractNum")
    an.set(qn("w:abstractNumId"), str(abs_id))
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal" if numbered else "bullet")
    lvl.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "%1." if numbered else "•")
    lvl.append(lvl_text)
    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    lvl.append(lvl_jc)
    ppr = OxmlElement("w:pPr")
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "680")
    ind.set(qn("w:hanging"), "340")
    ppr.append(ind)
    lvl.append(ppr)
    an.append(lvl)
    numbering.findall(qn("w:abstractNum"))[-1].addnext(an)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abs_ref = OxmlElement("w:abstractNumId")
    abs_ref.set(qn("w:val"), str(abs_id))
    num.append(abs_ref)
    numbering.findall(qn("w:num"))[-1].addnext(num)
    return num_id


def add_bullets(doc, items, numbered=False):
    num_id = _new_list_numbering(doc, numbered)
    for item in items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.25
        pPr = p._p.get_or_add_pPr()
        num_pr = OxmlElement("w:numPr")
        ilvl = OxmlElement("w:ilvl")
        ilvl.set(qn("w:val"), "0")
        num_pr.append(ilvl)
        num_id_el = OxmlElement("w:numId")
        num_id_el.set(qn("w:val"), str(num_id))
        num_pr.append(num_id_el)
        pPr.append(num_pr)
        add_runs(p, item, 11)


def _cell_para(cell, text, size=10.5, bold=False, color=DARK, align=None):
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.15
    if align is not None:
        p.alignment = align
    add_runs(p, text, size, bold=bold, color=color)
    return p


def add_table(doc, rows):
    header = [c.strip() for c in rows[0].strip("|").split("|")]
    body = [[c.strip() for c in r.strip("|").split("|")] for r in rows[2:]]
    n_cols = len(header)

    # 按内容长度分配列宽（中文按 2 字符计），最小 900 DXA
    lengths = []
    for j in range(n_cols):
        chars = len(header[j]) * 2
        for r in body:
            chars = max(chars, len(r[j]) * 2)
        lengths.append(max(chars, 8))
    total = sum(lengths)
    widths = [max(900, int(USABLE_DXA * c / total)) for c in lengths]
    widths[-1] = USABLE_DXA - sum(widths[:-1])

    table = doc.add_table(rows=1 + len(body), cols=n_cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False

    tbl = table._tbl
    tblPr = tbl.tblPr
    tbl_w = tblPr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tblPr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(USABLE_DXA))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tblPr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tblPr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    # 单元格边距（左右 108 DXA，上下 54 DXA）
    cell_mar = tblPr.find(qn("w:tblCellMar"))
    if cell_mar is None:
        cell_mar = OxmlElement("w:tblCellMar")
        tblPr.append(cell_mar)
    for side, val in (("top", "54"), ("left", "120"), ("bottom", "54"), ("right", "120")):
        el = cell_mar.find(qn("w:" + side))
        if el is None:
            el = OxmlElement("w:" + side)
            cell_mar.append(el)
        el.set(qn("w:w"), val)
        el.set(qn("w:type"), "dxa")

    # tblGrid
    grid = tbl.find(qn("w:tblGrid"))
    for gc in grid.findall(qn("w:gridCol")):
        grid.remove(gc)
    for w in widths:
        gc = OxmlElement("w:gridCol")
        gc.set(qn("w:w"), str(w))
        grid.append(gc)

    # 表头：品牌绿底白字 + 跨页重复
    for j, h in enumerate(header):
        cell = table.cell(0, j)
        cell.width = Cm(widths[j] * 2.54 / 1440)
        _cell_para(cell, h, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF), align=WD_ALIGN_PARAGRAPH.CENTER)
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), BRAND_GREEN)
        cell._tc.get_or_add_tcPr().append(shd)
    trPr = table.rows[0]._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    trPr.append(tbl_header)

    for i, row in enumerate(body, 1):
        for j, val in enumerate(row):
            cell = table.cell(i, j)
            cell.width = Cm(widths[j] * 2.54 / 1440)
            _cell_para(cell, val, align=WD_ALIGN_PARAGRAPH.CENTER if len(val) <= 12 else None)
            if i % 2 == 0:
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:fill"), LIGHT_GREEN)
                cell._tc.get_or_add_tcPr().append(shd)

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)
    spacer_run = spacer.add_run("")
    spacer_run.font.size = Pt(2)


def add_cover(doc):
    for _ in range(6):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("杏坛智备")
    set_run_font(run, 44, True, RGBColor.from_string(BRAND_GREEN), east="微软雅黑")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("乡村教师 AI 备课助手 · 商业计划书")
    set_run_font(run, 20, True, DARK, east="微软雅黑")
    for _ in range(2):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("中国国际大学生创新大赛（原“互联网+”）2027 · 青年红色筑梦之旅赛道 · 创意组")
    set_run_font(run, 12, False, GRAY, east="微软雅黑")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("版本 V1.2（镀金版）· 2026-08")
    set_run_font(run, 13, True, RGBColor.from_string(GOLD), east="微软雅黑")
    doc.add_page_break()


def add_footer(doc):
    footer = doc.sections[0].footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("杏坛智备 · 商业计划书 V1.2 · 第 ")
    set_run_font(run, 9, False, GRAY, east="宋体")
    run = p.add_run("")
    set_run_font(run, 9, False, GRAY, east="宋体")
    fld1 = OxmlElement("w:fldChar")
    fld1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld2 = OxmlElement("w:fldChar")
    fld2.set(qn("w:fldCharType"), "end")
    run._r.append(fld1)
    run._r.append(instr)
    run._r.append(fld2)
    run = p.add_run(" 页")
    set_run_font(run, 9, False, GRAY, east="宋体")


def main():
    lines = MD.read_text(encoding="utf-8").splitlines()
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Cm(21.0)
    sec.page_height = Cm(29.7)
    sec.top_margin = sec.bottom_margin = Cm(2.2)
    sec.left_margin = sec.right_margin = Cm(2.4)
    add_cover(doc)
    add_footer(doc)

    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        if not line.strip():
            i += 1
            continue
        if line.startswith("# "):
            add_heading(doc, line[2:].strip(), 1)
        elif line.startswith("## "):
            add_heading(doc, line[3:].strip(), 2)
        elif line.startswith("### "):
            add_heading(doc, line[4:].strip(), 3)
        elif line.startswith("#### "):
            add_heading(doc, line[5:].strip(), 4)
        elif line.startswith("> "):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.5)
            p.paragraph_format.space_after = Pt(4)
            add_runs(p, line[2:].strip(), 10.5, color=GRAY)
        elif re.match(r"^\|.+\|$", line):
            rows = [line]
            while i + 1 < len(lines) and re.match(r"^\|.+\|$", lines[i + 1].strip()):
                i += 1
                rows.append(lines[i].strip())
            if len(rows) >= 3 and re.match(r"^\|[\s:|-]+\|$", rows[1]):
                add_table(doc, rows)
            else:
                for r in rows:
                    add_para(doc, r)
        elif line.strip() in ("---", "***", "___"):
            pass
        elif re.match(r"^\d+\.\s+", line):
            items = [re.sub(r"^\d+\.\s+", "", line)]
            while i + 1 < len(lines) and re.match(r"^\d+\.\s+", lines[i + 1].strip()):
                i += 1
                items.append(re.sub(r"^\d+\.\s+", "", lines[i].strip()))
            add_bullets(doc, items, numbered=True)
        elif line.startswith("- "):
            items = [line[2:].strip()]
            while i + 1 < len(lines) and lines[i + 1].strip().startswith("- "):
                i += 1
                items.append(lines[i].strip()[2:].strip())
            add_bullets(doc, items)
        else:
            add_para(doc, line)
        i += 1

    doc.save(OUT)
    print(f"OK -> {OUT}")


if __name__ == "__main__":
    sys.exit(main())

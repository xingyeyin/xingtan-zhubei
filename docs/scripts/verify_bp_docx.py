# -*- coding: utf-8 -*-
import glob
from docx import Document

path = glob.glob(r"E:\比赛\互联网+\商业计划书-杏坛智备-V1.2.docx")[0]
d = Document(path)
text = "\n".join(p.text for p in d.paragraphs)
for t in d.tables:
    for row in t.rows:
        text += "\n" + " | ".join(c.text for c in row.cells)
checks = [
    "75 册", "476 个单元", "三情景敏感性分析", "乐观", "毛利率",
    "法律与合规声明", "0.3~0.8 元", "杏坛智备", "V1.2", "课标（9 科）",
    "2026-08-07", "让每一节乡村课都备得起、备得好",
]
all_ok = True
for c in checks:
    ok = c in text
    all_ok = all_ok and ok
    print(("OK  " if ok else "MISS"), c)
print("total chars:", len(text))
print("PASS" if all_ok else "FAIL")
